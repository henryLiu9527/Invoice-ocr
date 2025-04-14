import os
import json
import logging
import pandas as pd
from docx import Document
from config import RESULTS_FOLDER

logger = logging.getLogger(__name__)

class ResultExporter:
    """OCR结果导出器，支持多种格式导出"""
    
    def __init__(self):
        """初始化导出器"""
        logger.info("Initializing ResultExporter")
        
        # 确保结果目录存在
        os.makedirs(RESULTS_FOLDER, exist_ok=True)
        # 确保JSON结果目录存在
        os.makedirs(os.path.join(RESULTS_FOLDER, 'json'), exist_ok=True)
    
    def _get_output_path(self, original_filename, format_name):
        """
        生成输出文件路径
        
        Args:
            original_filename: 原始文件名
            format_name: 输出格式名称
            
        Returns:
            输出文件的完整路径
        """
        # 获取不带扩展名的原始文件名
        basename = os.path.splitext(os.path.basename(original_filename))[0]
        # 文件名安全处理
        safe_name = "".join([c for c in basename if c.isalnum() or c in (' ', '_', '-')]).rstrip()
        
        # 生成输出文件名
        timestamp = pd.Timestamp.now().strftime('%Y%m%d%H%M%S')
        output_filename = f"{safe_name}_{timestamp}.{format_name.lower()}"
        
        # 如果是JSON格式，放在json子目录中
        if format_name.lower() == 'json':
            return os.path.join(RESULTS_FOLDER, 'json', output_filename)
        else:
            return os.path.join(RESULTS_FOLDER, output_filename)
    
    def _extract_text_from_result(self, ocr_result):
        """
        从OCR结果中提取文本内容
        
        Args:
            ocr_result: OCR引擎返回的结果
            
        Returns:
            提取的文本列表
        """
        lines = []
        
        # 检查结果格式
        if not ocr_result.get('success', False):
            logger.warning("Cannot extract text: OCR result indicates failure")
            return ["OCR Processing Failed"]
        
        result = ocr_result.get('result', {})
        invoice_type = ocr_result.get('invoice_type', 'Auto')
        
        # 记录完整的结果结构，帮助调试
        logger.debug(f"OCR result structure: {json.dumps(result, ensure_ascii=False)[:500]}...")
        
        # 检查是否为空结果
        words_result_empty = False
        if 'words_result' in result:
            words_result = result.get('words_result', [])
            words_result_num = result.get('words_result_num', 0)
            if len(words_result) == 0 or words_result_num == 0:
                words_result_empty = True
                logger.warning("OCR result contains empty words_result array")
        
        # 检查发票数据提取 - 处理words_result中的特殊结构
        has_vat_invoice_data = False
        if 'words_result' in result and len(result['words_result']) > 0:
            # 检查words_result中是否包含result字段的vat_invoice数据
            for word_item in result['words_result']:
                if isinstance(word_item, dict) and 'result' in word_item and isinstance(word_item['result'], dict):
                    vat_data = word_item.get('result', {})
                    has_vat_invoice_data = True
                    logger.info(f"Found vat_invoice data in words_result")
                    lines.append("== 增值税发票识别结果 ==")
                    
                    # 1. 基本信息区域
                    lines.append("\n【基本信息】")
                    
                    # 发票类型和名称
                    invoice_type_org = self._extract_field_value(vat_data, 'InvoiceTypeOrg')
                    invoice_type_val = self._extract_field_value(vat_data, 'InvoiceType')
                    if invoice_type_org:
                        lines.append(f"发票名称: {invoice_type_org}")
                    if invoice_type_val:
                        lines.append(f"发票类型: {invoice_type_val}")
                    
                    # 发票代码和号码
                    invoice_code = self._extract_field_value(vat_data, 'InvoiceCode')
                    invoice_num = self._extract_field_value(vat_data, 'InvoiceNum')
                    if invoice_code:
                        lines.append(f"发票代码: {invoice_code}")
                    if invoice_num:
                        lines.append(f"发票号码: {invoice_num}")
                    
                    # 开票日期
                    invoice_date = self._extract_field_value(vat_data, 'InvoiceDate')
                    if invoice_date:
                        lines.append(f"开票日期: {invoice_date}")
                    
                    # 校验码
                    check_code = self._extract_field_value(vat_data, 'CheckCode')
                    if check_code:
                        lines.append(f"校验码: {check_code}")
                    
                    # 机器编号
                    machine_code = self._extract_field_value(vat_data, 'MachineCode')
                    if machine_code:
                        lines.append(f"机器编号: {machine_code}")
                    
                    # 2. 销售方信息区域
                    lines.append("\n【销售方信息】")
                    seller_name = self._extract_field_value(vat_data, 'SellerName')
                    seller_register_num = self._extract_field_value(vat_data, 'SellerRegisterNum')
                    seller_address = self._extract_field_value(vat_data, 'SellerAddress')
                    seller_bank = self._extract_field_value(vat_data, 'SellerBank')
                    
                    if seller_name:
                        lines.append(f"名称: {seller_name}")
                    if seller_register_num:
                        lines.append(f"纳税人识别号: {seller_register_num}")
                    if seller_address:
                        lines.append(f"地址、电话: {seller_address}")
                    if seller_bank:
                        lines.append(f"开户行及账号: {seller_bank}")
                    
                    # 3. 购买方信息区域
                    lines.append("\n【购买方信息】")
                    purchaser_name = self._extract_field_value(vat_data, 'PurchaserName')
                    purchaser_register_num = self._extract_field_value(vat_data, 'PurchaserRegisterNum')
                    purchaser_address = self._extract_field_value(vat_data, 'PurchaserAddress')
                    purchaser_bank = self._extract_field_value(vat_data, 'PurchaserBank')
                    
                    if purchaser_name:
                        lines.append(f"名称: {purchaser_name}")
                    if purchaser_register_num:
                        lines.append(f"纳税人识别号: {purchaser_register_num}")
                    if purchaser_address:
                        lines.append(f"地址、电话: {purchaser_address}")
                    if purchaser_bank:
                        lines.append(f"开户行及账号: {purchaser_bank}")
                    
                    # 4. 商品信息区域
                    lines.append("\n【商品信息】")
                    commodity_names = self._extract_field_list(vat_data, 'CommodityName')
                    commodity_prices = self._extract_field_list(vat_data, 'CommodityPrice')
                    commodity_nums = self._extract_field_list(vat_data, 'CommodityNum')
                    commodity_units = self._extract_field_list(vat_data, 'CommodityUnit')
                    commodity_tax_rates = self._extract_field_list(vat_data, 'CommodityTaxRate')
                    commodity_amounts = self._extract_field_list(vat_data, 'CommodityAmount')
                    commodity_taxes = self._extract_field_list(vat_data, 'CommodityTax')
                    
                    # 构建商品表格
                    if commodity_names:
                        lines.append("序号  商品名称                单价    数量  单位  税率     金额      税额")
                        lines.append("-------------------------------------------------------------------")
                        for i in range(len(commodity_names)):
                            name = commodity_names[i] if i < len(commodity_names) else ""
                            price = commodity_prices[i] if i < len(commodity_prices) else ""
                            num = commodity_nums[i] if i < len(commodity_nums) else ""
                            unit = commodity_units[i] if i < len(commodity_units) else ""
                            tax_rate = commodity_tax_rates[i] if i < len(commodity_tax_rates) else ""
                            amount = commodity_amounts[i] if i < len(commodity_amounts) else ""
                            tax = commodity_taxes[i] if i < len(commodity_taxes) else ""
                            
                            # 格式化为表格行
                            row = f"{i+1:<4} {name:<20} {price:<6} {num:<4} {unit:<4} {tax_rate:<6} {amount:<8} {tax}"
                            lines.append(row)
                    
                    # 5. 合计金额区域
                    lines.append("\n【金额信息】")
                    total_amount = self._extract_field_value(vat_data, 'TotalAmount')
                    total_tax = self._extract_field_value(vat_data, 'TotalTax')
                    amount_in_figuers = self._extract_field_value(vat_data, 'AmountInFiguers')
                    amount_in_words = self._extract_field_value(vat_data, 'AmountInWords')
                    
                    if total_amount:
                        lines.append(f"合计金额: {total_amount}")
                    if total_tax:
                        lines.append(f"合计税额: {total_tax}")
                    if amount_in_figuers:
                        lines.append(f"价税合计(小写): {amount_in_figuers}")
                    if amount_in_words:
                        lines.append(f"价税合计(大写): {amount_in_words}")
                    
                    # 6. 备注和其他信息
                    lines.append("\n【其他信息】")
                    remarks = self._extract_field_value(vat_data, 'Remarks')
                    payee = self._extract_field_value(vat_data, 'Payee')
                    checker = self._extract_field_value(vat_data, 'Checker')
                    note_drawer = self._extract_field_value(vat_data, 'NoteDrawer')
                    
                    if remarks:
                        lines.append(f"备注: {remarks}")
                    if payee:
                        lines.append(f"收款人: {payee}")
                    if checker:
                        lines.append(f"复核: {checker}")
                    if note_drawer:
                        lines.append(f"开票人: {note_drawer}")
                    
                    break
                
                # 如果在words_result项中找到了type字段为vat_invoice的数据
                elif isinstance(word_item, dict) and 'type' in word_item and word_item.get('type') == 'vat_invoice':
                    has_vat_invoice_data = True
                    logger.info(f"Found direct vat_invoice data in words_result")
                    
                    # 从原始JSON中提取结构化数据
                    if 'result' in word_item:
                        vat_data = word_item['result']
                        if isinstance(vat_data, dict):
                            lines.append("== 增值税发票识别结果 ==")
                            
                            # 提取各个字段并格式化
                            # 基本信息区域
                            lines.append("\n【基本信息】")
                            
                            # 发票类型和名称
                            for field, display in [
                                ('InvoiceTypeOrg', '发票名称'),
                                ('InvoiceType', '发票类型'),
                                ('InvoiceCode', '发票代码'),
                                ('InvoiceNum', '发票号码'),
                                ('InvoiceDate', '开票日期')
                            ]:
                                value = self._extract_field_value(vat_data, field)
                                if value:
                                    lines.append(f"{display}: {value}")
                                    
                            # 同样的方式处理其他区域...            
                    else:
                        # 没有找到result字段，简单展示
                        lines.append("== 增值税发票识别结果 ==")
                        for key, value in word_item.items():
                            if key not in ['type', 'probability', 'left', 'top', 'width', 'height']:
                                if isinstance(value, dict):
                                    value_str = json.dumps(value, ensure_ascii=False)
                                    if len(value_str) > 50:
                                        value_str = value_str[:47] + "..."
                                    lines.append(f"{key}: {value_str}")
                                else:
                                    lines.append(f"{key}: {value}")
                    
                    break
                    
        # 处理智能财务票据识别结果
        if (invoice_type == 'MultipleInvoice' or 'multiple_invoice' in result) and not has_vat_invoice_data:
            lines.append("== 智能财务票据识别结果 ==")
            
            # 获取multiple_invoice字段，可能在不同的层级
            multi_data = None
            if 'multiple_invoice' in result:
                multi_data = result.get('multiple_invoice', {})
            
            # 如果没有multiple_invoice字段，可能整个结果就是multiple_invoice结果
            if multi_data is None or not multi_data:
                logger.warning("MultipleInvoice type but no valid multiple_invoice data found")
                # 尝试直接使用结果
                multi_data = result
            
            # 输出完整的multi_data结构到日志，帮助调试
            logger.debug(f"MultipleInvoice data structure: {json.dumps(multi_data, ensure_ascii=False)[:500]}...")
            
            # 获取票据类型
            invoice_type_detected = None
            if isinstance(multi_data.get('invoice_type'), dict):
                invoice_type_detected = multi_data.get('invoice_type', {}).get('text', 'Unknown')
            elif isinstance(multi_data.get('invoice_type'), str):
                invoice_type_detected = multi_data.get('invoice_type')
            
            if invoice_type_detected:
                lines.append(f"检测到的票据类型: {invoice_type_detected}")
            else:
                lines.append("未检测到票据类型")
            
            lines.append("------------------------------")
            
            # 处理票据详细信息
            invoice_detail = multi_data.get('details', [])
            if isinstance(invoice_detail, list) and invoice_detail:
                for item in invoice_detail:
                    if isinstance(item, dict):
                        key = item.get('key', '')
                        value = item.get('value', '')
                        lines.append(f"{key}: {value}")
            else:
                logger.warning("No valid details found in MultipleInvoice data")
            
            # 处理票据金额
            if 'money' in multi_data:
                lines.append(f"金额: {multi_data['money']}")
            
            # 处理票据其他字段
            for field in ['tax', 'date', 'publisher', 'buyer', 'seller']:
                if field in multi_data:
                    lines.append(f"{field.capitalize()}: {multi_data[field]}")
            
            # 如果没有任何有效内容，检查是否有words_result
            if len(lines) <= 3:  # 只有标题和分隔线
                logger.warning("No valid content extracted from MultipleInvoice result, checking for words_result")
                if 'words_result' in result and not words_result_empty:
                    lines.append("识别到的文本内容:")
                    for item in result['words_result']:
                        if isinstance(item, dict) and 'words' in item:
                            lines.append(item['words'])
                else:
                    lines.append("未检测到文本内容")
                    
        # 处理表格识别结果
        elif invoice_type == 'Form' and not has_vat_invoice_data:
            lines.append("== 智能表格识别结果 ==")
            
            # 添加调试信息，输出收到的完整结果结构
            logger.info(f"Processing Form result with keys: {list(result.keys())}")
            if 'tables_result' in result:
                logger.info(f"Found tables_result with {len(result['tables_result'])} tables")
            elif 'result' in result and 'tables_result' in result['result']:
                logger.info(f"Found tables_result in result with {len(result['result']['tables_result'])} tables")
            
            # 检查是否有表格数据
            form_data = None
            if 'tables_result' in result:
                form_data = result.get('tables_result', [])
                logger.info("Using tables_result directly")
            elif 'result' in result and 'tables_result' in result['result']:
                form_data = result['result'].get('tables_result', [])
                logger.info("Using tables_result from result object")
            elif 'result' in result and 'forms_result' in result['result']:
                form_data = result['result'].get('forms_result', [])
                logger.info("Using forms_result from result object")
            elif 'forms_result' in result:
                form_data = result.get('forms_result', [])
                logger.info("Using forms_result directly")
            elif 'form_result' in result:
                form_data = result.get('form_result', [])
                logger.info("Using form_result directly")
            # 检查是否有单表格数据
            elif 'table_html' in result or ('result' in result and 'table_html' in result['result']):
                # 如果有table_html字段，说明是单表格
                table_html = result.get('table_html', '')
                if not table_html and 'result' in result:
                    table_html = result['result'].get('table_html', '')
                if table_html:
                    logger.info("Found table_html data")
                    lines.append(f"\n【HTML表格】")
                    lines.append("表格已识别，请查看JSON结果获取完整HTML表格")
                    form_data = [{'type': 'html'}]  # 创建一个虚拟表格项，使后续代码不会认为没有表格
            # 检查表格单元格数据
            elif 'tables' in result or ('result' in result and 'tables' in result['result']):
                if 'tables' in result:
                    form_data = result.get('tables', [])
                    logger.info("Using tables directly")
                else:
                    form_data = result['result'].get('tables', [])
                    logger.info("Using tables from result object")
            elif 'cells' in result or ('result' in result and 'cells' in result['result']):
                # 如果有cells字段，说明是使用了单元格格式返回表格
                cells = []
                matrix = []
                if 'cells' in result:
                    cells = result.get('cells', [])
                    matrix = result.get('matrix', [])
                    logger.info(f"Using cells directly: {len(cells)} cells with matrix size {len(matrix)}")
                elif 'result' in result:
                    cells = result['result'].get('cells', [])
                    matrix = result['result'].get('matrix', [])
                    logger.info(f"Using cells from result object: {len(cells)} cells with matrix size {len(matrix)}")
                
                if cells and matrix:
                    # 构建表格数据
                    table_data = []
                    for row in matrix:
                        row_data = []
                        for cell_idx in row:
                            if isinstance(cell_idx, int) and cell_idx < len(cells):
                                cell = cells[cell_idx]
                                row_data.append({'text': cell.get('text', '')})
                            else:
                                row_data.append({'text': ''})
                        table_data.append(row_data)
                    
                    form_data = [{'body': table_data}]
                    logger.info(f"Created table data with {len(table_data)} rows")
            
            if form_data and isinstance(form_data, list):
                logger.info(f"Processing {len(form_data)} tables in form result")
                
                # 处理每个表格
                for table_index, table in enumerate(form_data):
                    lines.append(f"\n【表格 {table_index + 1}】")
                    
                    # 跳过HTML表格，因为我们无法在文本中很好地显示它
                    if isinstance(table, dict) and table.get('type') == 'html':
                        lines.append("  (HTML表格 - 请查看导出的JSON文件获取完整内容)")
                        continue
                    
                    # 获取表格内容 - 处理标准body/header格式
                    if isinstance(table, dict) and ('body' in table or 'cells' in table):
                        # 处理标准body结构
                        body = table.get('body', [])
                        header = table.get('header', [])
                        cell_layout = {}  # 用于保存单元格布局
                        
                        # 特别处理百度表格识别API的body结构
                        if not body and 'cells' in table:
                            # 百度表格识别API返回的cell_location格式
                            logger.info(f"Processing Baidu Table API cells format")
                            # 提取所有单元格并按行组织
                            row_cells = {}
                            cell_texts = {}
                            row_heights = {}
                            col_widths = {}
                            
                            # 首先提取所有单元格信息
                            for cell in table.get('cells', []):
                                if isinstance(cell, dict):
                                    row_start = cell.get('row_start', 0)
                                    row_end = cell.get('row_end', row_start)
                                    col_start = cell.get('col_start', 0)
                                    col_end = cell.get('col_end', col_start)
                                    words = cell.get('words', '')
                                    
                                    # 记录单元格布局信息(跨行跨列)
                                    for r in range(row_start, row_end + 1):
                                        if r not in row_cells:
                                            row_cells[r] = []
                                            row_heights[r] = 0
                                        
                                        for c in range(col_start, col_end + 1):
                                            if c not in col_widths:
                                                col_widths[c] = 0
                                            
                                            # 只在左上角单元格保存实际内容
                                            if r == row_start and c == col_start:
                                                cell_texts[(r, c)] = words
                                                # 计算理想的列宽（基于内容长度）
                                                text_len = len(str(words))
                                                if text_len > col_widths[c]:
                                                    col_widths[c] = min(text_len, 30)  # 限制最大宽度为30
                                            
                                            # 记录单元格占用情况
                                            row_cells[r].append(c)
                                            cell_layout[(r, c)] = (row_start, col_start)
                            
                            # 确定所有行和列
                            max_row = max(row_cells.keys()) if row_cells else 0
                            all_cols = set()
                            for cols in row_cells.values():
                                all_cols.update(cols)
                            max_col = max(all_cols) if all_cols else 0
                            
                            # 构建表格内容
                            formatted_rows = []
                            for r in range(max_row + 1):
                                if r not in row_cells:
                                    continue
                                
                                row_text = []
                                for c in range(max_col + 1):
                                    if (r, c) in cell_layout:
                                        orig_r, orig_c = cell_layout[(r, c)]
                                        if r == orig_r and c == orig_c:
                                            # 这是一个单元格的左上角，显示内容
                                            content = cell_texts.get((r, c), '')
                                            width = col_widths.get(c, 10)
                                            row_text.append(f"{content:{width}}")
                                        # 否则这是一个被合并的单元格，不显示内容
                                    else:
                                        # 填充空单元格
                                        width = col_widths.get(c, 10)
                                        row_text.append(' ' * width)
                                
                                formatted_rows.append(row_text)
                            
                            # 输出格式化的表格
                            if formatted_rows:
                                # 生成分隔行
                                sep_line = '+'
                                for c in range(max_col + 1):
                                    width = col_widths.get(c, 10)
                                    sep_line += '-' * width + '+'
                                
                                lines.append(sep_line)
                                for row_text in formatted_rows:
                                    lines.append('|' + '|'.join(row_text) + '|')
                                    lines.append(sep_line)
                            else:
                                lines.append("  (表格内容解析失败)")
                        else:
                            # 标准body结构处理
                            logger.info(f"Table {table_index+1} has header: {bool(header)} and body with {len(body) if isinstance(body, list) else 'unknown'} rows")
                            
                            # 如果有表头，则显示表头
                            if header and isinstance(header, list):
                                # 计算每列的最大宽度
                                col_widths = []
                                header_text = []
                                
                                for cell in header:
                                    # 获取单元格文本
                                    cell_text = ""
                                    if isinstance(cell, dict):
                                        if 'words' in cell:
                                            cell_text = cell['words']
                                        elif 'text' in cell:
                                            cell_text = cell['text']
                                    else:
                                        cell_text = str(cell)
                                    
                                    header_text.append(cell_text)
                                    # 更新列宽
                                    while len(col_widths) < len(header_text):
                                        col_widths.append(0)
                                    col_widths[len(header_text)-1] = max(col_widths[len(header_text)-1], min(len(cell_text), 30))
                                
                                # 生成表头行
                                header_row = "| "
                                for i, text in enumerate(header_text):
                                    width = max(col_widths[i], 5)  # 最小宽度5
                                    header_row += f"{text:{width}} | "
                                lines.append(header_row)
                                
                                # 生成分隔行
                                separator = "+-"
                                for width in col_widths:
                                    width = max(width, 5)  # 最小宽度5
                                    separator += "-" * width + "-+-"
                                lines.append(separator)
                            
                            # 显示表格内容
                            if body and isinstance(body, list):
                                for row_idx, row in enumerate(body):
                                    if isinstance(row, list):
                                        # 处理列表形式的行
                                        row_text = "| "
                                        for i, cell in enumerate(row):
                                            # 获取单元格文本
                                            cell_text = ""
                                            if isinstance(cell, dict):
                                                if 'words' in cell:
                                                    cell_text = cell['words']
                                                elif 'text' in cell:
                                                    cell_text = cell['text']
                                            else:
                                                cell_text = str(cell)
                                            
                                            # 确保col_widths有足够的元素
                                            while i >= len(col_widths):
                                                col_widths.append(0)
                                            
                                            # 更新并使用列宽
                                            width = max(col_widths[i], 5)  # 最小宽度5
                                            row_text += f"{cell_text:{width}} | "
                                        
                                        lines.append(row_text)
                                    # 百度表格识别API也可能返回行对象
                                    elif isinstance(row, dict):
                                        # 处理对象形式的行
                                        if 'words' in row:
                                            # 这是百度表格识别的行对象
                                            lines.append(f"| {row.get('words', '')} |")
                                        elif 'col_start' in row and 'col_end' in row:
                                            # 这是单元格对象
                                            col_span = row.get('col_end', 0) - row.get('col_start', 0) + 1
                                            words = row.get('words', '')
                                            lines.append(f"| {words} {'|' * col_span}")
                                        elif 'cells' in row:
                                            # 行包含多个单元格
                                            row_text = "| "
                                            for i, cell in enumerate(row['cells']):
                                                cell_text = cell.get('text', '')
                                                # 确保col_widths有足够的元素
                                                while i >= len(col_widths):
                                                    col_widths.append(0)
                                                
                                                # 更新并使用列宽
                                                width = max(col_widths[i], 5)  # 最小宽度5
                                                row_text += f"{cell_text:{width}} | "
                                            
                                            lines.append(row_text)
                            elif body and isinstance(body, dict):
                                # 处理字典形式的body
                                lines.append("  (表格结构复杂，请查看JSON结果获取完整内容)")
                            else:
                                lines.append("  (表格内容为空)")
                    else:
                        # 尝试处理直接的表格数据
                        if isinstance(table, list):
                            # 计算每列的最大宽度
                            col_widths = []
                            
                            # 第一遍：计算所有列宽
                            for row in table:
                                if isinstance(row, list):
                                    for i, cell in enumerate(row):
                                        cell_text = ""
                                        if isinstance(cell, str):
                                            cell_text = cell
                                        elif isinstance(cell, dict):
                                            if 'text' in cell:
                                                cell_text = cell['text']
                                            elif 'words' in cell:
                                                cell_text = cell['words']
                                        
                                        # 更新列宽
                                        while i >= len(col_widths):
                                            col_widths.append(0)
                                        col_widths[i] = max(col_widths[i], min(len(str(cell_text)), 30))
                            
                            # 生成分隔行
                            separator = "+-"
                            for width in col_widths:
                                width = max(width, 5)  # 最小宽度5
                                separator += "-" * width + "-+-"
                            
                            # 显示表格数据
                            lines.append(separator)
                            for row in table:
                                if isinstance(row, list):
                                    row_text = "| "
                                    for i, cell in enumerate(row):
                                        cell_text = ""
                                        if isinstance(cell, str):
                                            cell_text = cell
                                        elif isinstance(cell, dict):
                                            if 'text' in cell:
                                                cell_text = cell['text']
                                            elif 'words' in cell:
                                                cell_text = cell['words']
                                        
                                        # 使用列宽格式化
                                        width = max(col_widths[i], 5) if i < len(col_widths) else 10
                                        row_text += f"{cell_text:{width}} | "
                                    
                                    lines.append(row_text)
                                    lines.append(separator)
                                elif isinstance(row, dict) and 'words' in row:
                                    lines.append(f"| {row['words']} |")
                                    lines.append(separator)
                        else:
                            lines.append("  (无法解析的表格格式)")
            else:
                logger.warning("No valid table data found in Form result")
                
                # 如果没有表格数据，但有words_result，则展示文本内容
                if 'words_result' in result and not words_result_empty:
                    lines.append("识别到的文本内容:")
                    for item in result['words_result']:
                        if isinstance(item, dict) and 'words' in item:
                            lines.append(item['words'])
                else:
                    lines.append("未检测到表格内容，OCR结果可能不包含表格数据")
                    
                # 记录完整的结果结构，帮助调试
                logger.debug(f"Form result structure: {json.dumps(result, ensure_ascii=False)[:500]}...")
        
        # 处理通用高精度接口返回结果 (accurate)
        elif invoice_type == 'Accurate' and 'words_result' in result and not has_vat_invoice_data:
            lines.append("== 通用高精度识别结果 ==")
            words_result = result.get('words_result', [])
            
            if words_result_empty:
                lines.append("未检测到文本内容")
            else:
                for item in words_result:
                    if isinstance(item, dict):
                        words = item.get('words', '')
                        probability = item.get('probability', {}).get('average', 0)
                        probability_str = f" (置信度: {probability:.2f})" if probability else ""
                        lines.append(f"{words}{probability_str}")
                        
                        # 如果有位置信息，也可以添加
                        location = item.get('location', {})
                        if location:
                            loc_str = f"位置: 左={location.get('left', 0)}, 上={location.get('top', 0)}, 宽={location.get('width', 0)}, 高={location.get('height', 0)}"
                            lines.append(f"  {loc_str}")
        
        # 处理标准的words_result格式
        elif 'words_result' in result and not has_vat_invoice_data:
            words_result = result.get('words_result', [])
            
            if words_result_empty:
                lines.append("未检测到文本内容")
            else:
                for item in words_result:
                    if isinstance(item, dict) and 'words' in item:
                        lines.append(item['words'])
                    elif isinstance(item, str):
                        lines.append(item)
        
        # 处理特殊格式（增值税发票等）- 这里处理顶层vat_invoice
        if 'vat_invoice' in result and not has_vat_invoice_data:
            vat_data = result['vat_invoice']
            lines.append("== 增值税发票识别结果 ==")
            
            # 基本信息区域
            lines.append("\n【基本信息】")
            invoice_fields = [
                ('InvoiceTypeOrg', '发票名称'),
                ('InvoiceType', '发票类型'),
                ('InvoiceCode', '发票代码'),
                ('InvoiceNum', '发票号码'),
                ('InvoiceDate', '开票日期'),
                ('CheckCode', '校验码'),
                ('MachineCode', '机器编号')
            ]
            
            for field, display in invoice_fields:
                value = self._extract_field_value(vat_data, field)
                if value:
                    lines.append(f"{display}: {value}")
            
            # 销售方信息
            lines.append("\n【销售方信息】")
            seller_fields = [
                ('SellerName', '名称'),
                ('SellerRegisterNum', '纳税人识别号'),
                ('SellerAddress', '地址、电话'),
                ('SellerBank', '开户行及账号')
            ]
            
            for field, display in seller_fields:
                value = self._extract_field_value(vat_data, field)
                if value:
                    lines.append(f"{display}: {value}")
            
            # 购买方信息
            lines.append("\n【购买方信息】")
            purchaser_fields = [
                ('PurchaserName', '名称'),
                ('PurchaserRegisterNum', '纳税人识别号'),
                ('PurchaserAddress', '地址、电话'),
                ('PurchaserBank', '开户行及账号')
            ]
            
            for field, display in purchaser_fields:
                value = self._extract_field_value(vat_data, field)
                if value:
                    lines.append(f"{display}: {value}")
            
            # 商品信息
            lines.append("\n【商品信息】")
            if 'CommodityName' in vat_data:
                commodity_names = self._extract_field_list(vat_data, 'CommodityName')
                commodity_prices = self._extract_field_list(vat_data, 'CommodityPrice')
                commodity_nums = self._extract_field_list(vat_data, 'CommodityNum')
                commodity_units = self._extract_field_list(vat_data, 'CommodityUnit')
                commodity_tax_rates = self._extract_field_list(vat_data, 'CommodityTaxRate')
                commodity_amounts = self._extract_field_list(vat_data, 'CommodityAmount')
                commodity_taxes = self._extract_field_list(vat_data, 'CommodityTax')
                
                # 构建商品表格
                if commodity_names:
                    lines.append("序号  商品名称                单价    数量  单位  税率     金额      税额")
                    lines.append("-------------------------------------------------------------------")
                    for i in range(len(commodity_names)):
                        name = commodity_names[i] if i < len(commodity_names) else ""
                        price = commodity_prices[i] if i < len(commodity_prices) else ""
                        num = commodity_nums[i] if i < len(commodity_nums) else ""
                        unit = commodity_units[i] if i < len(commodity_units) else ""
                        tax_rate = commodity_tax_rates[i] if i < len(commodity_tax_rates) else ""
                        amount = commodity_amounts[i] if i < len(commodity_amounts) else ""
                        tax = commodity_taxes[i] if i < len(commodity_taxes) else ""
                        
                        # 格式化为表格行
                        row = f"{i+1:<4} {name:<20} {price:<6} {num:<4} {unit:<4} {tax_rate:<6} {amount:<8} {tax}"
                        lines.append(row)
            
            # 金额信息
            lines.append("\n【金额信息】")
            amount_fields = [
                ('TotalAmount', '合计金额'),
                ('TotalTax', '合计税额'),
                ('AmountInFiguers', '价税合计(小写)'),
                ('AmountInWords', '价税合计(大写)')
            ]
            
            for field, display in amount_fields:
                value = self._extract_field_value(vat_data, field)
                if value:
                    lines.append(f"{display}: {value}")
            
            # 其他信息
            lines.append("\n【其他信息】")
            other_fields = [
                ('Remarks', '备注'),
                ('Payee', '收款人'),
                ('Checker', '复核'),
                ('NoteDrawer', '开票人')
            ]
            
            for field, display in other_fields:
                value = self._extract_field_value(vat_data, field)
                if value:
                    lines.append(f"{display}: {value}")
        
        # 如果没有提取到任何文本
        if not lines:
            logger.warning("No text could be extracted from the OCR result, adding default message")
            lines.append("No text detected or unsupported result format")
            # 记录结果结构，帮助调试
            logger.debug(f"OCR result keys: {list(result.keys())}")
        
        return lines
    
    def _extract_field_value(self, data, field_name):
        """
        从字段中提取值
        
        Args:
            data: 包含字段的字典
            field_name: 字段名
            
        Returns:
            提取的值，如果不存在则返回空字符串
        """
        if field_name not in data:
            return ""
        
        field_data = data[field_name]
        
        # 如果是列表，取第一个元素
        if isinstance(field_data, list) and len(field_data) > 0:
            item = field_data[0]
            # 如果是字典且包含word字段
            if isinstance(item, dict) and 'word' in item:
                return item['word']
            # 否则直接返回该元素的字符串形式
            return str(item)
        
        # 如果是字典且包含word字段
        if isinstance(field_data, dict) and 'word' in field_data:
            return field_data['word']
        
        # 其他情况直接返回字符串形式
        return str(field_data)
    
    def _extract_field_list(self, data, field_name):
        """
        从字段中提取值列表
        
        Args:
            data: 包含字段的字典
            field_name: 字段名
            
        Returns:
            提取的值列表，如果不存在则返回空列表
        """
        result = []
        
        if field_name not in data:
            return result
        
        field_data = data[field_name]
        
        # 如果不是列表，直接返回空列表
        if not isinstance(field_data, list):
            return result
            
        # 处理列表
        for item in field_data:
            if isinstance(item, dict) and 'word' in item:
                result.append(item['word'])
            else:
                result.append(str(item))
        
        return result
    
    def export_to_json(self, ocr_result, original_filename):
        """
        将原始OCR结果保存为JSON格式
        
        Args:
            ocr_result: OCR引擎返回的结果
            original_filename: 原始文件名
            
        Returns:
            输出文件路径
        """
        logger.info(f"Saving original OCR result to JSON for: {os.path.basename(original_filename)}")
        
        try:
            output_path = self._get_output_path(original_filename, 'json')
            
            # 准备要保存的数据
            export_data = {
                'metadata': {
                    'original_file': os.path.basename(original_filename),
                    'ocr_engine': ocr_result.get('engine', 'unknown'),
                    'invoice_type': ocr_result.get('invoice_type', 'unknown'),
                    'export_time': pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')
                },
                'ocr_result': ocr_result
            }
            
            # 保存为JSON文件
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"JSON export successful: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to export to JSON: {str(e)}")
            raise
    
    def export_to_xlsx(self, ocr_result, original_filename):
        """
        将OCR结果导出为Excel格式
        
        Args:
            ocr_result: OCR引擎返回的结果
            original_filename: 原始文件名
            
        Returns:
            输出文件路径
        """
        logger.info(f"Exporting OCR result to XLSX for: {os.path.basename(original_filename)}")
        
        try:
            output_path = self._get_output_path(original_filename, 'xlsx')
            
            # 提取文本
            lines = self._extract_text_from_result(ocr_result)
            
            # 创建DataFrame
            df = pd.DataFrame(lines, columns=['Text'])
            
            # 添加元数据
            metadata = {
                'Original File': os.path.basename(original_filename),
                'OCR Engine': ocr_result.get('engine', 'unknown'),
                'Invoice Type': ocr_result.get('invoice_type', 'unknown'),
                'Export Time': pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            # 创建一个Excel写入器
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                # 写入OCR结果
                df.to_excel(writer, sheet_name='OCR Results', index=False)
                
                # 写入元数据
                pd.DataFrame([metadata]).T.to_excel(writer, sheet_name='Metadata')
            
            logger.info(f"XLSX export successful: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to export to XLSX: {str(e)}")
            raise
    
    def export_to_docx(self, ocr_result, original_filename):
        """
        将OCR结果导出为Word文档格式
        
        Args:
            ocr_result: OCR引擎返回的结果
            original_filename: 原始文件名
            
        Returns:
            输出文件路径
        """
        logger.info(f"Exporting OCR result to DOCX for: {os.path.basename(original_filename)}")
        
        try:
            output_path = self._get_output_path(original_filename, 'docx')
            
            # 提取文本
            lines = self._extract_text_from_result(ocr_result)
            
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            doc.add_heading('OCR Recognition Results', level=1)
            
            # 添加元数据
            doc.add_heading('Metadata', level=2)
            doc.add_paragraph(f"Original File: {os.path.basename(original_filename)}")
            doc.add_paragraph(f"OCR Engine: {ocr_result.get('engine', 'unknown')}")
            doc.add_paragraph(f"Invoice Type: {ocr_result.get('invoice_type', 'unknown')}")
            doc.add_paragraph(f"Export Time: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # 添加OCR结果
            doc.add_heading('Recognized Text', level=2)
            for line in lines:
                doc.add_paragraph(line)
            
            # 保存文档
            doc.save(output_path)
            
            logger.info(f"DOCX export successful: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to export to DOCX: {str(e)}")
            raise
    
    def export_to_txt(self, ocr_result, original_filename):
        """
        将OCR结果导出为文本格式
        
        Args:
            ocr_result: OCR引擎返回的结果
            original_filename: 原始文件名
            
        Returns:
            输出文件路径
        """
        logger.info(f"Exporting OCR result to TXT for: {os.path.basename(original_filename)}")
        
        try:
            output_path = self._get_output_path(original_filename, 'txt')
            
            # 提取文本
            lines = self._extract_text_from_result(ocr_result)
            
            # 写入文本文件
            with open(output_path, 'w', encoding='utf-8') as f:
                # 写入元数据
                f.write(f"# OCR Recognition Results\n")
                f.write(f"# Original File: {os.path.basename(original_filename)}\n")
                f.write(f"# OCR Engine: {ocr_result.get('engine', 'unknown')}\n")
                f.write(f"# Invoice Type: {ocr_result.get('invoice_type', 'unknown')}\n")
                f.write(f"# Export Time: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("#" + "-" * 50 + "\n\n")
                
                # 写入OCR结果
                for line in lines:
                    f.write(f"{line}\n")
            
            logger.info(f"TXT export successful: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to export to TXT: {str(e)}")
            raise
    
    def export(self, ocr_result, original_filename, format_name):
        """
        将OCR结果导出为指定格式
        
        Args:
            ocr_result: OCR引擎返回的结果
            original_filename: 原始文件名
            format_name: 导出格式（xlsx, docx, txt, json）
            
        Returns:
            输出文件路径
        """
        logger.info(f"Exporting OCR result to {format_name} for: {os.path.basename(original_filename)}")
        
        # 总是保存一份原始JSON，无论请求什么格式
        try:
            self.export_to_json(ocr_result, original_filename)
        except Exception as e:
            logger.error(f"Failed to save JSON backup: {str(e)}")
        
        format_name = format_name.lower()
        
        if format_name == 'xlsx':
            return self.export_to_xlsx(ocr_result, original_filename)
        elif format_name == 'docx':
            return self.export_to_docx(ocr_result, original_filename)
        elif format_name == 'txt':
            return self.export_to_txt(ocr_result, original_filename)
        elif format_name == 'json':
            return self.export_to_json(ocr_result, original_filename)
        else:
            logger.error(f"Unsupported export format: {format_name}")
            raise ValueError(f"Unsupported export format: {format_name}") 